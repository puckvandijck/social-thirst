"""
Streamlit app for TikTok content idea generation (V3).

V3 introduces live trend scraping via Apify and content generation via OpenAI.

Features:
 - Uses Apify's TikTok Trends Scraper (clockworks/tiktok-trends-scraper) to
   fetch trending hashtags and songs for the selected region and industry. This
   requires the user to set the environment variable `APIFY_API_TOKEN` with
   their Apify API token. If no token is provided, the app falls back to the
   stub trends defined in V2.
 - Generates content ideas using OpenAI's API (GPT-5 or later). The user must
   set `OPENAI_API_KEY` to their OpenAI API key. A fallback heuristic generator
   is provided if no key is found.
 - Integrates vectorization of captions using TF-IDF to provide context to the
   language model and to compute similarity between trends and existing posts.
 - Respects the brand book guidelines (placeholder parser) and includes them in
   the prompt sent to the model.

Before running this app, install the necessary dependencies:

```
pip install streamlit pandas numpy python-docx scikit-learn openai apify-client
```

Set up environment variables in a `.env` file or your shell:

```
export OPENAI_API_KEY="sk-..."
export APIFY_API_TOKEN="apify-api-token"
```

Then run the app with:

```
streamlit run streamlit_app_v3.py
```

Replace the placeholders with your actual API tokens. The app will attempt to
fetch live trends from Apify; if the token is missing, stub data will be used.
"""

import os
import regex as re
from collections import Counter
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st
# Map Streamlit secrets to environment variables (for cloud & local parity)
import os
try:
    for _k in ("OPENAI_API_KEY", "APIFY_API_TOKEN"):
        if _k in st.secrets and st.secrets[_k] and not os.environ.get(_k):
            os.environ[_k] = st.secrets[_k]
except Exception:
    pass
from docx import Document
from docx.shared import Pt

# Optional external libraries
try:
    from sklearn.feature_extraction.text import TfidfVectorizer  # type: ignore
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

try:
    import openai  # type: ignore
    OPENAI_AVAILABLE = True
except Exception:
    OPENAI_AVAILABLE = False

try:
    from apify_client import ApifyClient  # type: ignore
    APIFY_AVAILABLE = True
except Exception:
    APIFY_AVAILABLE = False

# -------------------------------
# i18n / Language support
# -------------------------------
import streamlit as st

# Keep the code readable by using keys you can re-use across the app
TRANSLATIONS = {
    "nl": {
        "app_title": "Content Tool",
        "sidebar_settings": "Instellingen",
        "choose_region": "Kies regio",
        "regions": ["Nederland", "België", "Duitsland", "Verenigde Staten"],
        "time_range_label": "Selecteer tijdsperiode voor trends",
        "time_ranges": ["7 dagen", "30 dagen", "120 dagen"],
        "brand_label": "Selecteer merk",
        "brand_help": "Kies het merk waarvoor je contentideeën wilt genereren.",
        "product_push_title": "💄 Product of range om te pushen (optioneel)",
        "product_push_desc": ("Kies een platform om te pushen in het dropdown-menu of vul zelf handmatig één of meerdere producten in. "
                              "Indien je voor platform kiest, wordt minstens één product uit dit platform op een natuurlijke manier in het idee verwerkt. "
                              "Laat beide leeg als je geen specifieke producten wilt highlighten."),
        "generate_button": "Genereer ideeën",
        "err_upload_one": "Upload minstens één analytics-bestand.",
        "err_no_videos": "Geen videobestand gevonden in de uploads.",
        "ideas_header": "Gegenereerde contentideeën",
        "live_hashtags": "Live trending hashtags (Scraped)",
        "live_sounds": "Live trending sounds (Scraped)",
        "main_trend_focus": "Hoofdtrend om op te focussen",
        "performing_tags": "Presterende hashtags uit je data",
        "upload_brandbook": "Upload Masterbrand Data",
        "apify_missing": "Apify niet beschikbaar in deze omgeving.",
        "apify_no_token": "Geen Apify token gevonden. Stel APIFY_API_TOKEN in.",
        "openai_key_missing": "OPENAI_API_KEY is niet ingesteld. Kan brand book niet analyseren.",
        "openai_call_error": "❌ Fout bij het aanroepen van OpenAI: {err}",
        "docx_title": "Idee {i}: {title}",
        "docx_trend": "Trendfocus: ",
        "docx_why_now": "Waarom actueel: ",
        "docx_hook": "Hook: ",
        "docx_script": "Script:",
        "docx_tips": "Filmtips: ",
        "docx_hashtags": "Hashtags: ",
        "docx_sound": "Sound: ",
        "docx_sound_ok": " (Commercieel toegestaan: {ok})",
        "lang_label": "Taal / Language",
        "lang_options": {"Nederlands": "nl", "English": "en"},
        # New UI and message keys for Dutch
        "cohort_label": "Cohort om op te focussen (optioneel)",
        "cohort_help": "Bijv. 'Romance Seekers'",
        "main_trend_label": "Trend om op te focussen (optioneel)",
        "main_trend_help": "Bijv. 'group7', 'soft girl aesthetic', 'de-influencing', enz.",
        "low_budget_label": "Low‑budget ideeën",
        "low_budget_help": "ideeën bevatten slechts één locatie en acteur",
        "select_files_label": "Selecteer één of meerdere CSV/XLSX bestanden",
        "upload_brand_book_label": "Upload je brand book (PDF, optioneel)",
        "no_preference": "Geen voorkeur",
        "platform_label": "Platform (optioneel)",
        "manual_products_label": "Handmatig product(en) toevoegen (optioneel)",
        "product_name": "Naam",
        "product_description": "Beschrijving",
        "no_main_trend_message": "Geen hoofdtrend ingevuld — live trends van Apify worden automatisch gebruikt.",
        "fetching_trends_message": "Bezig met ophalen van TikTok trends voor",
        "average_success_score": "Gemiddelde successcore",
        "artist": "Artiest",
        "commercial_allowed": "Commercieel toegestaan",
        "sound": "Sound",
        "unknown_artist": "Onbekende artiest",
        "yes": "Ja",
        "no": "Nee",
        "download_csv": "📊 Download als CSV",
        "download_docx": "📄 Download als DOCX",
        "caption": "Bijschrift",
        "docx_caption": "Bijschrift: ",
        "no_live_hashtags_warning": "Geen live hashtags gevonden (controleer filters of land).",
        "no_live_sounds_warning": "Geen live sounds gevonden (controleer filters of land).",
    },
    "en": {
        "app_title": "Content Tool",
        "sidebar_settings": "Settings",
        "choose_region": "Choose region",
        "regions": ["Netherlands", "Belgium", "Germany", "United States"],
        "time_range_label": "Select trend time window",
        "time_ranges": ["7 days", "30 days", "120 days"],
        "brand_label": "Select brand",
        "brand_help": "Pick the brand you want ideas for.",
        "product_push_title": "💄 Product/range to promote (optional)",
        "product_push_desc": ("Choose a platform in the dropdown or manually add one or more products. "
                              "If you choose a platform, at least one product from that platform is woven naturally into each idea. "
                              "Leave both empty if you don’t want to highlight specific products."),
        "generate_button": "Generate ideas",
        "err_upload_one": "Please upload at least one analytics file.",
        "err_no_videos": "No video sheet found in the uploads.",
        "ideas_header": "Generated content ideas",
        "live_hashtags": "Live trending hashtags (Scraped)",
        "live_sounds": "Live trending sounds (Scraped)",
        "main_trend_focus": "Main trend to focus on",
        "performing_tags": "Top-performing hashtags from your data",
        "upload_brandbook": "Upload Masterbrand Data",
        "apify_missing": "Apify is not available in this environment.",
        "apify_no_token": "No Apify token found. Please set APIFY_API_TOKEN.",
        "openai_key_missing": "OPENAI_API_KEY is not set. Cannot analyze brand book.",
        "openai_call_error": "❌ Error calling OpenAI: {err}",
        "docx_title": "Idea {i}: {title}",
        "docx_trend": "Trend focus: ",
        "docx_why_now": "Why it’s timely: ",
        "docx_hook": "Hook: ",
        "docx_script": "Script:",
        "docx_tips": "Filming tips: ",
        "docx_hashtags": "Hashtags: ",
        "docx_sound": "Sound: ",
        "docx_sound_ok": " (Commercial OK: {ok})",
        "lang_label": "Taal / Language",
        "lang_options": {"Nederlands": "nl", "English": "en"},
        # New UI and message keys for English
        "cohort_label": "Cohort to focus on (optional)",
        "cohort_help": "e.g. 'Romance Seekers'",
        "main_trend_label": "Trend to focus on (optional)",
        "main_trend_help": "e.g. 'group7', 'soft girl aesthetic', 'de-influencing', etc.",
        "low_budget_label": "Low‑budget ideas",
        "low_budget_help": "ideas involve only one location and actor",
        "select_files_label": "Select one or more CSV/XLSX files",
        "upload_brand_book_label": "Upload your brand book (PDF, optional)",
        "no_preference": "No preference",
        "platform_label": "Platform (optional)",
        "manual_products_label": "Add product(s) manually (optional)",
        "product_name": "Name",
        "product_description": "Description",
        "no_main_trend_message": "No main trend entered — live trends from Apify will be used automatically.",
        "fetching_trends_message": "Retrieving TikTok trends for",
        "average_success_score": "Average success score",
        "artist": "Artist",
        "commercial_allowed": "Commercially allowed",
        "sound": "Sound",
        "unknown_artist": "Unknown artist",
        "yes": "Yes",
        "no": "No",
        "download_csv": "📊 Download as CSV",
        "download_docx": "📄 Download as DOCX",
        "caption": "Caption",
        "docx_caption": "Caption: ",
        "no_live_hashtags_warning": "No live hashtags found (check filters or country).",
        "no_live_sounds_warning": "No live sounds found (check filters or country).",
    },
}

def get_lang():
    # default to Dutch
    if "lang" not in st.session_state:
        st.session_state["lang"] = "nl"
    return st.session_state["lang"]

def t(key):
    lang = get_lang()
    # fallback to nl → key string itself if missing
    return TRANSLATIONS.get(lang, TRANSLATIONS["nl"]).get(key, TRANSLATIONS["nl"].get(key, key))

def t_format(key, **kwargs):
    return t(key).format(**kwargs)



# -------------------------------
# Helper functions (reused from V2)
# -------------------------------

import regex as re
import unicodedata

def is_latin_hashtag(tag: str) -> bool:
    """
    Return True if the hashtag text is primarily Latin-based
    (Dutch, English, German, etc.) and not Cyrillic, Arabic, or Asian.
    """

    if not tag:
        return False

    # Remove hashtag symbol and trim
    tag = tag.lstrip("#").strip()

    # Normalize and remove invisible Unicode marks
    tag = unicodedata.normalize("NFKC", tag)
    tag = re.sub(r"[\u200b-\u200f\u202a-\u202e\u2060-\u206f]", "", tag)

    # Reject if non-Latin scripts are present
    if re.search(r'[\p{IsCyrillic}\p{IsArabic}\p{IsHan}\p{IsHiragana}\p{IsKatakana}]', tag):
        return False

    # Keep if contains at least one Latin letter
    return bool(re.search(r'[A-Za-zÀ-ÿ]', tag))

def filter_trending_hashtags(tags: List[str], niche: str) -> List[str]:
    """
    Filter trending hashtags for relevancy to the given niche and remove banned or controversial topics.
    Only tags that relate to the niche (e.g. hair care, body care, beauty) or generic event-related words
    are kept. Tags containing political, election or sports-related keywords are excluded.
    """
    banned_keywords = [
        "politic", "politiek", "election", "vote", "voetbal", "soccer", "football",
        "war", "conflict", "controversy", "verkiezing", "verkiezingen"
    ]
    niche_keywords_map = {
        "Hair Care": [
            "hair", "haar", "hairstyle", "haarstijl", "haircare", "haarverzorging",
            "beauty", "glow", "party", "feest", "festival", "dance", "event"
        ],
        "Body Care": [
            "body", "skin", "huid", "bodycare", "body care", "lotion", "moisturizer",
            "wellness", "beauty"
        ],
        "Beauty": [
            "beauty", "glow", "makeup", "style", "fashion", "party", "fest", "festival", "event"
        ],
    }
    filtered: List[str] = []
    for tag in tags:
        tag_name = str(tag).lower()
        # Skip banned or controversial topics
        if any(bad in tag_name for bad in banned_keywords):
            continue
        allowed_keywords = niche_keywords_map.get(niche, [])
        # Keep tag if it contains any niche keyword
        if any(keyword in tag_name for keyword in allowed_keywords):
            filtered.append(tag)
        else:
            # Keep general event/festival/party tags if not explicitly banned
            if any(word in tag_name for word in ["event", "festival", "party", "feest", "fest"]):
                filtered.append(tag)
    return filtered


def detect_file_type(df: pd.DataFrame) -> str:
    cols = set(c.lower().strip() for c in df.columns)
    if {"video title", "video views", "video link"}.issubset(cols):
        return "video"
    if {"profile views", "video views", "reached audience"}.issubset(cols):
        return "overview"
    if {"new followers", "total followers", "engaged audience"}.issubset(cols):
        return "audience"
    return "unknown"


def load_analytics(files: List) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    video_df = pd.DataFrame()
    overview_df = pd.DataFrame()
    audience_df = pd.DataFrame()
    for uploaded_file in files:
        if uploaded_file is None:
            continue
        try:
            df = pd.read_csv(uploaded_file)
        except Exception:
            try:
                df = pd.read_excel(uploaded_file)
            except Exception:
                st.warning(f"Unable to read {uploaded_file.name} as CSV or Excel.")
                continue
        ftype = detect_file_type(df)
        if ftype == "video":
            video_df = df
        elif ftype == "overview":
            overview_df = df
        elif ftype == "audience":
            audience_df = df
    return video_df, overview_df, audience_df


def normalise_video_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [c.strip() for c in df.columns]
    if "Post time" in df.columns:
        df["Post time"] = pd.to_datetime(df["Post time"], errors="coerce")
        df["Date"] = df["Post time"].dt.date
    for col in ["Video views", "Likes", "Comments", "Shares", "Add to Favorites"]:
        if col not in df.columns:
            df[col] = 0
        df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)
    df["Engagement"] = df[["Likes", "Comments", "Shares", "Add to Favorites"]].sum(axis=1)
    df["Engagement rate"] = df["Engagement"] / df["Video views"].replace(0, np.nan)
    df["Engagement rate"] = df["Engagement rate"].fillna(0)
    if not df["Engagement rate"].empty:
        max_rate = df["Engagement rate"].max()
        if max_rate > 0:
            df["Success score"] = df["Engagement rate"] / max_rate
        else:
            df["Success score"] = 0
    else:
        df["Success score"] = 0
    return df


def extract_performing_hashtags(video_df: pd.DataFrame, top_n: int = 10) -> List[Tuple[str, float]]:
    if video_df.empty or "Success score" not in video_df.columns:
        return []
    hashtag_pattern = re.compile(r"#(\w+)")
    hashtag_scores: Dict[str, List[float]] = {}
    for _, row in video_df.iterrows():
        title = row.get("Video title", "")
        success = float(row.get("Success score", 0))
        if isinstance(title, str):
            tags = hashtag_pattern.findall(title)
            for tag in tags:
                tag_l = tag.lower()
                hashtag_scores.setdefault(tag_l, []).append(success)
    averages: Dict[str, float] = {
        tag: (sum(scores) / len(scores) if scores else 0)
        for tag, scores in hashtag_scores.items()
    }
    sorted_tags = sorted(averages.items(), key=lambda x: x[1], reverse=True)
    return sorted_tags[:top_n]


def parse_brandbook_placeholder() -> Dict[str, List[str]]:
    return {
        "tone": "neutraal, vriendelijk, deskundig",
        "avoid": ["garantie", "geneest", "wondermiddel"],
        "must": ["Andrélon", "natuurlijk", "haarverzorging"],
    }


def vectorize_captions(video_df: pd.DataFrame) -> Tuple[Optional[TfidfVectorizer], Optional[np.ndarray]]:
    if not SKLEARN_AVAILABLE:
        return None, None
    titles = [str(t) if not pd.isna(t) else "" for t in video_df.get("Video title", [])]
    if not titles:
        return None, None
    vectorizer = TfidfVectorizer(max_features=200)
    matrix = vectorizer.fit_transform(titles)
    return vectorizer, matrix

from typing import List, Dict, Any
import json

def analyse_trends_meanings(
    client,
    trends: List[str],
    brand_name: str,
    niche: str,
    max_trends: int = 12,
    model: str = "gpt-5.1-mini",
) -> Dict[str, str]:
    """
    Calls the model once to produce short Dutch explanations for each trend.
    Returns a dict {trend: explanation}. Safe for empty/noisy inputs.
    """
    trends = [t for t in (trends or []) if isinstance(t, str) and t.strip()]
    if not trends:
        return {}

    # Deduplicate while preserving order; limit to avoid token bloat
    seen = set()
    clean = []
    for t in trends:
        key = t.strip()
        if key not in seen:
            seen.add(key)
            clean.append(key)
        if len(clean) >= max_trends:
            break

    # Ask for strict JSON to keep parsing robust
    user_prompt = f"""
Je bent een Nederlandse TikTok-trendanalist voor de {niche}-niche (merk: {brand_name}).
Geef per trend een KORTE, praktische uitleg in maximaal 2 zinnen:
- Wat houdt de trend in?
- Hoe manifesteert deze zich in TikTok-content (stijl/format/doelgroep)?

Antwoord ALLEEN als strikt JSON-object waarbij de keys exact de aangeleverde trendnamen zijn, bijv.:
{{
  "clean girl aesthetic": "Korte uitleg...",
  "get ready with me": "Korte uitleg..."
}}

Trends:
{json.dumps(clean, ensure_ascii=False)}
"""

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "Je bent een beknopte, nauwkeurige assisent die uitsluitend valide JSON terugstuurt."},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
        )
        raw = resp.choices[0].message.content.strip()
        # Try to extract JSON even if the model wrapped it in code fences
        if raw.startswith("```"):
            raw = raw.strip("`")
            # remove possible language hints like ```json
            first_newline = raw.find("\n")
            if first_newline != -1:
                raw = raw[first_newline+1:].strip()
        parsed = json.loads(raw)
        # Keep only requested trends; coerce to str
        out: Dict[str, str] = {}
        for t in clean:
            val = parsed.get(t, "")
            if isinstance(val, str):
                out[t] = val.strip()
            else:
                out[t] = str(val)
        return out
    except Exception:
        # Fail-soft: return empty dict on any parsing/network issues
        return {}

# --- Helper: Clean escaped newlines and tabs from GPT output ---
def clean_output_text(text: str) -> str:
    """Replace escaped newlines and stray backslashes with real newlines/tabs."""
    if not isinstance(text, str):
        return text
    return text.replace("\\n", "\n").replace("\\t", "\t").strip()


# helper product to push dropdown

# ---- Load Andrélon platform catalog (cached) ----
@st.cache_data(show_spinner=False)
def load_platform_catalog(xlsx_path: str | None = None) -> pd.DataFrame:
    """Load the Andrelon platformen Excel, either uploaded or local fallback."""
    import pandas as pd
    import io

    candidates = [xlsx_path, "Andrelon platformen.xlsx", "andrelon_platformen.xlsx"]
    for path in candidates:
        if path is None:
            continue
        try:
            if hasattr(path, "read"):  # uploaded file-like object
                df = pd.read_excel(path, engine="openpyxl")
            else:
                df = pd.read_excel(path, engine="openpyxl")
            if not df.empty:
                return df
        except Exception as e:
            print(f"⚠️ Kon bestand niet lezen: {path} — {e}")
            continue

    return pd.DataFrame(columns=["Range", "SKU", "Description"])


# -------------------------------
# Live trend scraping via Apify (robust parser)
# -------------------------------
def fetch_apify_trends(
    region: str,
    industry: str = "Beauty & Personal Care",
    results_per_page: int = 50,
    time_range: str = "7"   # ✅ new parameter
) -> Tuple[List[str], List[Dict[str, str]]]:

    """
    Returns:
      - hashtags: List[str]
      - sounds: List[{"soundName": str, "artist": str, "isCommercial": bool}]
    """
    if not APIFY_AVAILABLE:
        st.warning(t("apify_missing"))
        return [], []

    token = os.environ.get("APIFY_API_TOKEN")
    if not token:
        st.warning(t("apify_no_token"))
        return [], []

    region_to_code = {
        "Nederland": "NL",
        "België": "BE",
        "Duitsland": "DE",
        "Verenigde Staten": "US",
    }
    country_code = region_to_code.get(region, "NL")

    client = ApifyClient(token)
    run_input = {
        "adsCountryCode": country_code,
        "adsApprovedForBusinessUse": True,   # keep True to align with brand-safe music
        "adsScrapeHashtags": True,
        "adsScrapeSounds": True,
        "adsScrapeCreators": False,
        "adsScrapeVideos": False,
        "adsRankType": "popular",
        "adsTimeRange": time_range,   # ✅ dynamic timeframe from user selection,
        "resultsPerPage": results_per_page,
    }

    try:
        run = client.actor("clockworks/tiktok-trends-scraper").call(run_input=run_input, timeout_secs=1200)
    except Exception as e:
        st.error(f"Error running Apify actor: {e}")
        return [], []

    dataset_id = run.get("defaultDatasetId")
    if not dataset_id:
        st.error("No dataset returned from Apify run.")
        return [], []

    hashtags: List[str] = []
    sounds: List[Dict[str, str]] = []

    try:
        for item in client.dataset(dataset_id).iterate_items():
            # Some actors return dicts, never strings — but be defensive:
            if not isinstance(item, dict):
                continue

            itype = str(item.get("type", "")).lower()

            # --- Hashtags ---
            if "hashtag" in itype:
                name = item.get("hashtagName") or item.get("title") or item.get("name")
                if name and name not in hashtags:
                    hashtags.append(name)
                continue

            # --- Sounds (sometimes 'sound', sometimes 'music') ---
            if ("sound" in itype) or ("music" in itype):
                sound_name = item.get("soundName") or item.get("musicName") or item.get("title") or item.get("name")
                artist = item.get("authorName") or item.get("artistName") or item.get("creatorName") or "Onbekend"
                commercial = item.get("isCommercial")
                if commercial is None:
                    commercial = item.get("commercial_ok", item.get("commercial", False))
                if sound_name:
                    sounds.append({
                        "soundName": sound_name,
                        "artist": artist,
                        "isCommercial": bool(commercial),
                    })
                continue

            # Optional: heuristic fallback (rare)
            url = item.get("url", "")
            if url and "/tag/" in url:
                name = item.get("title") or item.get("name")
                if name and name not in hashtags:
                    hashtags.append(name)

    except Exception as e:
        st.error(f"Error parsing Apify dataset: {e}")
        return [], []

    if not hashtags:
        st.warning(t("no_live_hashtags_warning"))
    if not sounds:
        st.warning(t("no_live_sounds_warning"))

    return hashtags, sounds



# -------------------------------
# Idea generation with OpenAI (improved V2 — trend-weighted & diverse)
# -------------------------------

from openai import OpenAI
import os, json
from datetime import datetime
import pandas as pd
import streamlit as st

def generate_ideas_openai(
    video_df: pd.DataFrame,
    trending_tags: List[str],
    sounds: List[Dict[str, str]],
    brand_cfg: Dict[str, List[str]],
    n_ideas: int = 5,
    niche: str = "Beauty",
    product_push_table: pd.DataFrame | None = None,
    cohort_focus: str = "",
    main_trend_focus: str = "",
    low_budget: bool = False,
    lang: str = "nl",
) -> pd.DataFrame:
    """Generate diverse, trend-weighted TikTok ideas using OpenAI GPT-5."""
    trend_analyses = {}  # ✅ ensures variable exists in all cases
    
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.warning(t("openai_key_missing"))

    client = OpenAI(api_key=api_key)

    # --- Context preparation ---
    # Use only the top 3 videos to prevent over-indexing on historical performance
    top_videos = video_df.sort_values(by="Success score", ascending=False).head(3)
    video_summaries = [
        {
            "title": str(row.get("Video title", "")),
            "success_score": float(row.get("Success score", 0)),
            "engagement_rate": float(row.get("Engagement rate", 0)),
        }
        for _, row in top_videos.iterrows()
    ]

    tone = brand_cfg.get("tone", "")
    avoid_words = ", ".join(brand_cfg.get("avoid", []))
    must_words = ", ".join(brand_cfg.get("must", []))
    today = datetime.now().strftime("%d %B %Y")

    # --- Determine which trends to include in the prompt + analyse meanings ---
    trend_section_text = "### Trends\nGeen trendinformatie beschikbaar."

    if main_trend_focus and main_trend_focus.strip():
        # Manual main trend → ignore Apify & analytics trends
        focus_raw = main_trend_focus.strip()
        focus_nohash = focus_raw.lstrip("#")
        # Show BOTH the hashtag and raw term to anchor the exact token in the model
        top_hashtags = [f"#{focus_nohash}", focus_raw]
        sounds_for_prompt = []
        trends_to_analyse = [focus_raw]

        brand_name_for_ai = brand_cfg.get("brand", "onbekend")
        trend_analyses = analyse_trends_meanings(
            client=client,
            trends=trends_to_analyse,
            brand_name=brand_name_for_ai,
            niche=niche,
            max_trends=12,
            model="gpt-5.1-mini",
        )
    
    analyses_markdown = ""
    if trend_analyses:
        bullets = []
        for trend_name in trends_to_analyse:
            expl = trend_analyses.get(trend_name, "").strip()
            bullets.append(f"- **{trend_name}** — {expl}" if expl else f"- **{trend_name}**")
        analyses_markdown = "\n\n### Korte trendanalyses\n" + "\n".join(bullets)

    # --- FIXED STRUCTURE ---
    if main_trend_focus and main_trend_focus.strip():
        trend_section_text = (
            "### Hoofdtrend om op te focussen\n"
            f"{focus_raw}\n"
            "### Verplichte hashtag/term\n"
            f"#{focus_nohash} (en/of '{focus_raw}')\n"
            "### Instructie\n"
            f"Behandel deze hoofdtrend als bewezen relevant voor {brand_cfg.get('brand','het merk')}. "
            "Integreer deze hoofdtrend in elk idee en gebruik de exacte tag/term in de hashtag-lijst."
            f"{analyses_markdown}"
        )

    else:
        # Build list of allowed trending hashtags from 'trending_tags' parameter
        allowed_tags = []
        for tag_item in (trending_tags or []):
            if not tag_item:
                continue
            tag = tag_item if tag_item.startswith("#") else f"#{tag_item}"
            if tag not in allowed_tags:
                allowed_tags.append(tag)

        if allowed_tags:
            bullet_list = "\n".join(f"- {t}" for t in allowed_tags[:10])
            trend_section_text = (
                "### Toegestane trending hashtags (EXACTE strings)\n"
                f"{bullet_list}\n"
                "### Instructie\n"
                "Kies per idee precies één van de bovenstaande hashtags als het veld 'trend_used'. "
                "Neem die tag ongewijzigd (exacte string) op in 'hashtags'. "
                "Pas de tag NIET aan: geen extra woorden, geen meervoud, geen prefix/suffix, geen varianten."
            )
        else:
            trend_section_text = "### Trends\nGeen trendinformatie beschikbaar."

        
    # --- Build dynamic products section ---
    products_section = ""

    # Get chosen platform name, if available
    # Retrieve platform choice, defaulting to the localized "no preference"
    platform_choice = st.session_state.get("platform_choice", t("no_preference"))

    if product_push_table is not None and not product_push_table.empty:
        # Convert manual/auto-filled table to readable lines
        products_text = "\n".join(
            f"- {row.get('Naam') or row.get('Productnaam', 'Onbekend product')}: {row.get('Beschrijving', '').strip()}"
            for _, row in product_push_table.iterrows()
            if str(row.get('Naam') or row.get('Productnaam')).strip()
        )
        products_section = (
            f"Er is een productfocus op de volgende items:\n{products_text}\n"
            "Zorg dat minstens één van deze producten op een natuurlijke manier voorkomt in elk idee."
        )
    elif platform_choice and platform_choice != t("no_preference"):
        products_section = (
            f"Het merk wil momenteel het platform **{platform_choice}** pushen. "
            "Zorg dat minstens één product uit dit platform op een natuurlijke manier voorkomt in elk idee."
        )
    else:
        products_section = "Er is geen specifieke product- of platformfocus opgegeven."


    # Force output language for titles, scripts, explanations, hashtags
    _language_name = "Nederlands" if lang == "nl" else "English"
    language_instruction = (
        f"IMPORTANT: Write ALL titles, scripts, short explanations, hashtags and any free text in {_language_name}. "
        f"If {_language_name} is English, avoid Dutch words and translate trend/context as needed."
    )


    #### prompt ####

    prompt = f"""
    {language_instruction}

    Je bent een creatieve TikTok-strateeg gespecialiseerd in Gen Z-content en merkactivaties voor het merk {brand_cfg.get('brand', 'onbekend')} in de {niche}-niche. 
    Gebruik de volgende informatie om {n_ideas} nieuwe, actuele en creatieve TikTok-video-ideeën te bedenken.
    

    ### Informatie

    # Merktoon
    {tone or "Professioneel maar speels"}
    Moet bevatten: {must_words or "geen specifieke richtlijnen"}
    Vermijd: {avoid_words or "geen specifieke verboden woorden"}

    # Product/range om te verwerken in content idee
    {products_section}

    # Toppresterende video's (context, niet dominant)
    {video_summaries}

    # Trending hashtags (dominant)
    {trend_section_text}


    ### Opdracht
    1. Analyseer welke trends het meest relevant zijn voor {brand_cfg.get('brand', 'het merk')}.    
    2. Bedenk vervolgens {n_ideas} creatieve TikTok-ideeën die inspelen op deze trends. De trends zijn tijdelijk en belangrijker dan historische data — gebruik ze actief in je ideeën.    
    3. Zorg dat de ideeën variëren in type (educatief, humor, transformatie, lifestyle, storytelling, enz.), en vermijd reclamejargon.    
    4. {(        
      "Gebruik uitsluitend de opgegeven hoofdtrend. "        
      "Voeg de exacte term toe aan de hashtaglijst van ieder idee: "        
      f"#{focus_nohash} (en/of '{focus_raw}'). Negeer overige trends."        
      if main_trend_focus and main_trend_focus.strip()        
      else "Integreer minimaal één actuele trend (hashtag of sound) per idee en leg de koppeling kort uit."    
    )}    
    5. Gebruik de merktoon, maar geef ruimte aan creatieve vrijheid.    
    6. {(        
        f"Focus op het cohort **{cohort_focus.strip()}**. " if cohort_focus else ""    
    )}{(        
        "Integreer minstens één van de genoemde producten of het gekozen platform in elk idee."         
        if products_section and "geen specifieke" not in products_section.lower() else ""    
    )}    
    7. Beschrijf per idee een creatief “Script” in 4–6 korte shots (SHOT 1, SHOT 2, …) met voice‑over, on‑screen text of handelingen.     
       Zorg dat de shots logisch op elkaar volgen binnen één herkenbare setting of een natuurlijk verloop. Geen abrupte wissels zonder context, zorg dat de flow duidelijk is.     
       Zorg dat het merk en product natuurlijk in het verhaal passen.    
    8. Vermijd politieke of controversiële onderwerpen en zorg dat de trends logisch passen bij de niche.     
       Bijvoorbeeld: voor haarverzorging zijn feesten en evenementen relevant (zoals festivals), maar sportwedstrijden     
       of politieke kwesties niet.    
    9. Geef voor elk idee een korte "caption" van maximaal één zin die het idee samenvat.
       {"10. **Low budget:** elk script mag maximaal 1 filmlocatie en 1 acteur bevatten." if low_budget else ""}


    BELANGRIJKE REGELS VOOR TRENDS EN HASHTAGS:
    - Gebruik een trending hashtag alleen als deze duidelijk aansluit bij het merk of de niche.
    - Kies per idee één trending hashtag en zet die exact (ongewijzigd, inclusief #) in 'trend_used'.
    - Neem diezelfde hashtag ook letterlijk op in 'hashtags'.
    - 'hashtags' bevat 3–5 items die allemaal met # beginnen.
    - 'caption' bevat maximaal één zin en beschrijft de essentie van het idee.
    - Geef uitsluitend valide JSON terug met de velden: title, hook, script, filming_tips, hashtags, sound, commercial_ok, short_explanation, trend_used, caption.


    """

    # --- GPT call ---
    try:
        # Define the system message based on language
        system_prompt = (
            "Je bent een ervaren TikTok-marketeer en trendanalist."
            if lang == "nl"
            else "You are an experienced TikTok marketer and trend analyst."
        )

        response = client.chat.completions.create(
            model="gpt-5",  # or "gpt-4o" if preferred
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            temperature=1.0,  # increase creativity and variation
        )

    except Exception as e:
        st.error(t_format("openai_call_error", err=e))


    # continue normally if no error
    records = []
    allowed_tags_set = set()
    # Reconstruct the allowed set so we can validate, based on what the prompt showed:
    if main_trend_focus and main_trend_focus.strip():
        focus_raw = main_trend_focus.strip()
        focus_tag = focus_raw if focus_raw.startswith("#") else f"#{focus_raw.lstrip('#')}"
        allowed_tags_set.add(focus_tag.lower())
    else:
        for tag_item in (trending_tags or []):
            if not tag_item:
                continue
            tag = tag_item if tag_item.startswith("#") else f"#{tag_item}"
            allowed_tags_set.add(tag.lower())

    # --- Parse model reply safely ---
    try:
        reply = response.choices[0].message.content.strip()
        if reply.startswith("```"):
            reply = reply.strip("`").split("json", 1)[-1].strip()
        ideas_list = json.loads(reply)
    except Exception as e:
        st.warning(f"⚠️ Fout bij het lezen van modeloutput: {e}")
        ideas_list = []

    # continue
    for idea in ideas_list[:n_ideas]:
        title = idea.get("title", "")
        hook = idea.get("hook", "")
        script = idea.get("script", "")

        # filming tips → string
        filming_tips_raw = idea.get("filming_tips", "")
        if isinstance(filming_tips_raw, list):
            filming_tips = "; ".join(map(str, filming_tips_raw))
        else:
            filming_tips = str(filming_tips_raw or "")

        # hashtags → list normalization
        hashtags_raw = idea.get("hashtags", [])
        if isinstance(hashtags_raw, str):
            # split on commas/spaces, keep entries with a leading #
            parts = [h.strip() for h in re.split(r"[,\s]+", hashtags_raw) if h.strip()]
            hashtags = [h if h.startswith("#") else f"#{h.lstrip('#')}" for h in parts]
        elif isinstance(hashtags_raw, list):
            hashtags = [h if str(h).startswith("#") else f"#{str(h).lstrip('#')}" for h in hashtags_raw]
        else:
            hashtags = []

        # trend_used → force exact presence in hashtags if it's valid/allowed
        trend_used = idea.get("trend_used", "")
        if trend_used:
            trend_used = trend_used if trend_used.startswith("#") else f"#{trend_used.lstrip('#')}"
            # If we have an allowed set, prefer only exact allowed tags
            if allowed_tags_set and trend_used.lower() not in allowed_tags_set:
                # if model invented/modified a tag, try to fallback to first allowed tag
                trend_used = next(iter(allowed_tags_set)).lower()
                # ensure it has '#'
                if not str(trend_used).startswith("#"):
                    trend_used = f"#{str(trend_used).lstrip('#')}"
            # ensure trend_used is included in hashtags (exact string)
            if trend_used not in hashtags:
                hashtags = [trend_used] + [h for h in hashtags if h.lower() != trend_used.lower()]
        else:
            # No trend_used provided: if allowed set exists, insert the first allowed tag
            if allowed_tags_set:
                trend_used = next(iter(allowed_tags_set))
                if not str(trend_used).startswith("#"):
                    trend_used = f"#{str(trend_used).lstrip('#')}"
                if trend_used not in hashtags:
                    hashtags = [trend_used] + hashtags

        # sound & flags
        sound = idea.get("sound", "")
        # Localize commercial OK flag
        commercial_ok = t("yes") if idea.get("commercial_ok", True) else t("no")
        short_explanation = idea.get("short_explanation", "")
        caption = idea.get("caption", "")

        records.append({
            "Titel": title,
            "Hook": hook,
            "Script": script,
            "Filming tips": filming_tips,
            "Hashtags": ", ".join(dict.fromkeys(hashtags)),  # keep order, dedupe
            "Sound": sound,
            "Commercial OK": commercial_ok,
            "Waarom actueel": short_explanation,
            "Caption": caption,
            "Trend": trend_used or "",  # expose for later export/use
        })

    return pd.DataFrame(records)


# -------------------------------
# Brand book PDF parser (improved robust version)
# -------------------------------
import PyPDF2
import regex as re
import json
import ast
from openai import OpenAI

def clean_pdf_text(text: str) -> str:
    """Normalize extracted PDF text for better GPT comprehension."""
    text = re.sub(r"\s{2,}", " ", text)  # collapse multiple spaces
    text = re.sub(r"\n{2,}", "\n", text)  # reduce multiple newlines
    text = re.sub(r"[^a-zA-Z0-9.,;:!?()&%€$'\-\n ]", "", text)  # strip weird symbols
    text = re.sub(r"(?i)page\s*\d+", "", text)  # remove page markers
    text = text.strip()
    return text

def parse_brandbook(brandbook_file, brand_name: str = "Onbekend") -> Dict[str, List[str]]:
    """Extract tone, must/avoid words, and brand essence from a brand book PDF using GPT."""
    if not brandbook_file:
        # No upload — just return empty/default config
        return {"tone": "", "must": [], "avoid": []}

    # --- Step 1: Read and clean PDF text ---
    pdf_reader = PyPDF2.PdfReader(brandbook_file)
    text = ""
    for page in pdf_reader.pages:
        try:
            text += page.extract_text() or ""
        except Exception:
            continue

    cleaned_text = clean_pdf_text(text)

    if len(cleaned_text) < 500:
        st.warning("PDF bevat weinig leesbare tekst — mogelijk is het document grotendeels beeldmateriaal.")
        return {"tone": "", "must": [], "avoid": []}

    # --- Step 2: Analyze using GPT ---
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.warning("OPENAI_API_KEY is niet ingesteld. Kan brand book niet analyseren.")
        return {"tone": "", "must": [], "avoid": []}

    client = OpenAI(api_key=api_key)

    prompt = f"""
    Je bent een Nederlandse merkstrateeg. Analyseer de onderstaande tekst uit een merkboek van {brand_name}.
    Beschrijf:
    1. De toon van stem (max 40 woorden)
    2. 3–5 kernwoorden of thema’s die de merkidentiteit definiëren ("must")
    3. 3–5 woorden of concepten die de merktoon zouden verstoren ("avoid")

    Geef het antwoord als geldige JSON met velden: tone, must, avoid.

    MERKBOEK TEKST:
    {cleaned_text}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-5",  # or "gpt-4o" depending on your setup
            messages=[
                {"role": "system", "content": "Je bent een ervaren Nederlandse merkstrateeg."},
                {"role": "user", "content": prompt},
            ],
        )
        reply = response.choices[0].message.content.strip()

        try:
            cfg = json.loads(reply)
        except Exception:
            cfg = ast.literal_eval(reply)

        # Validate JSON structure
        if not isinstance(cfg, dict):
            raise ValueError("Invalid JSON structure returned by GPT")

        # Ensure required keys exist
        for key in ["tone", "must", "avoid"]:
            if key not in cfg:
                cfg[key] = [] if key != "tone" else ""

        return cfg

    except Exception as e:
        st.error(f"Fout bij het analyseren van brand book: {e}")
        return {"tone": "", "must": [], "avoid": []}

# -------------------------------
# Streamlit UI
# -------------------------------
def main():
    st.set_page_config(
        layout="wide"
    )

    st.title(t("app_title"))


    # --- Custom light-purple theme styling ---
    st.markdown("""
        <style>
            /* Main app background */
            .stApp {
                background-color: #FFFFFF !important;  /* white lavender */
            }
    
            /* Sidebar background and text */
            section[data-testid="stSidebar"] {
                background-color: #EDE7F6;  /* light pastel purple */
                color: #2C1A4D;  /* deep plum text */
            }

            /* Sidebar headings and labels */
            section[data-testid="stSidebar"] h1,
            section[data-testid="stSidebar"] h2,
            section[data-testid="stSidebar"] h3,
            section[data-testid="stSidebar"] p,
            section[data-testid="stSidebar"] label,
            section[data-testid="stSidebar"] span {
                color: #2C1A4D !important;
            }

            /* Sidebar buttons */
            section[data-testid="stSidebar"] button {
                background-color: #6A1B9A !important;  /* strong purple */
                color: #FFFFFF !important;
                border-radius: 8px !important;
            }

            /* Sidebar button hover */
            section[data-testid="stSidebar"] button:hover {
                background-color: #4A0072 !important;
                color: #FFFFFF !important;
            }

            /* Main area buttons */
            button[kind="primary"] {
                background-color: #6A1B9A !important;
                color: #FFFFFF !important;
                border-radius: 8px !important;
            }

            /* Subtle card effect for main content blocks */
            div[data-testid="stVerticalBlock"] > div {
                background-color: #FFFFFF;
                padding: 1.2rem;
                border-radius: 12px;
                box-shadow: 0 2px 8px rgba(106, 27, 154, 0.08); /* soft purple shadow */
                margin-bottom: 1rem;
            }

            /* General text color */
            div[data-testid="stMarkdownContainer"] p {
                color: #2C1A4D;
                font-size: 16px;
            }

            /* Divider line */
            hr {
                border: 1px solid #E0D7F6;
            }

            /* Headers */
            h1, h2, h3 {
                color: #4A0072;
            }
        </style>
    """, unsafe_allow_html=True)


    # --- Sidebar controls ---

    with st.sidebar:
        # Language picker at the top of the sidebar
        _lang_label = t("lang_label")
        _lang_options = list(TRANSLATIONS["nl"]["lang_options"].keys())  # show human labels
        _default_label = "Nederlands" if get_lang() == "nl" else "English"
        chosen = st.selectbox(_lang_label, _lang_options, index=_lang_options.index(_default_label))
        st.session_state["lang"] = TRANSLATIONS["nl"]["lang_options"][chosen]

        st.header(t("sidebar_settings"))
 
        _region_options = t("regions")
        region = st.selectbox(
            t("choose_region"),
            _region_options,
            index=0,
            help= t("choose_region")  # short/help keeps consistent; customize if you want separate help text
        )


            # --- Time range selector for trends ---
        time_range = st.selectbox(
            t("time_range_label"),
            t("time_ranges"),
            index=0
        )
        # Normalize to numeric days regardless of language
        _time_map_nl = {"7 dagen": "7", "30 dagen": "30", "120 dagen": "120"}
        _time_map_en = {"7 days": "7", "30 days": "30", "120 days": "120"}
        selected_time_range = (_time_map_nl.get(time_range) or _time_map_en.get(time_range) or "7")

        # brand select
        brand = st.selectbox(
            t("brand_label"),
            ["Andrelon", "Dove", "Vaseline"],
            index=0,
            help=t("brand_help")
        )        


        # Backend mapping of brand → niche
        brand_to_niche = {
            "Andrelon": "Hair Care",
            "Dove": "Body Care",
            "Vaseline": "Body Care",
        }
        niche = brand_to_niche.get(brand, "Beauty")

        ## cohort selection ##
        cohort_focus = st.text_input(
            t("cohort_label"),
            value="",
            help=t("cohort_help")
        )


        # --- Optional main trend focus ---
        main_trend_focus = st.text_input(
            t("main_trend_label"),
            value="",
            help=t("main_trend_help")
        )

        # NEW: Low-budget toggle
        low_budget = st.checkbox(
            t("low_budget_label"),
            value=False,  # default unchecked
            help=t("low_budget_help")  # localized help
        )

        # (Optional) keep it around for later use elsewhere too
        st.session_state["low_budget"] = low_budget

    # upload docs
    st.subheader(t("upload_brandbook"))

    # --- Analytics file uploader (optional) ---
    analytics_files = st.file_uploader(
        t("select_files_label"),
        type=["csv", "xlsx"],
        accept_multiple_files=True,
        help=t("select_files_help"),
    )

    analytics_dataframes = []
    if analytics_files:
        for uploaded_file in analytics_files:
            df = read_uploaded_file(uploaded_file)
            if df is not None:
                analytics_dataframes.append(df)


    # --- Brand book uploader (optional) ---
    brandbook_file = st.file_uploader(
        t("upload_brand_book_label"),
        type=["pdf"],
        help=t("upload_brand_book_help"),
    )


    # --- 🧴 Product of range / platform keuze (optioneel) ---
    st.markdown("### " + t("product_push_title"))
    st.write(t("product_push_desc"))


    # --- Load the fixed Excel catalog silently ---
    catalog_df = load_platform_catalog()

    # --- Platform dropdown ---
    # Build platform options with localized "no preference"
    platform_options = [t("no_preference")]
    if not catalog_df.empty:
        unique_platforms = sorted(catalog_df["Range"].dropna().unique().tolist())
        platform_options.extend(unique_platforms)

    platform_choice = st.selectbox(
        t("platform_label"),
        platform_options,
        key="platform_choice"
    )


    # --- Manual product entry ---
    # Manual product entry section (localized header)
    st.markdown("#### " + t("manual_products_label"))

    # Prepare a default DataFrame for product entry with internal column names
    default_products = pd.DataFrame([{"Naam": "", "Beschrijving": ""}])
    # Create a copy for display with localized column headers
    display_products = default_products.rename(
        columns={"Naam": t("product_name"), "Beschrijving": t("product_description")}
    )
    # Let the user edit the products using the localized display; then rename back to internal names
    manual_products_df_display = st.data_editor(
        display_products,
        num_rows="dynamic",
        width=1000,
        key="manual_push_products"
    )
    # Convert back to internal column names for later processing
    manual_products_df = manual_products_df_display.rename(
        columns={t("product_name"): "Naam", t("product_description"): "Beschrijving"}
    )

    # --- Merge data from dropdown & manual input ---
    if platform_choice != t("no_preference") and not catalog_df.empty:
        auto_products_df = catalog_df[catalog_df["Range"] == platform_choice][["SKU", "Description"]].copy()
        auto_products_df.columns = ["Naam", "Beschrijving"]
    else:
        auto_products_df = pd.DataFrame(columns=["Naam", "Beschrijving"])

    merged_products_df = pd.concat(
        [auto_products_df, manual_products_df],
        ignore_index=True
    ).fillna("")

    # Clean up
    merged_products_df = merged_products_df.astype(str)
    merged_products_df["Naam"] = merged_products_df["Naam"].str.strip()
    merged_products_df["Beschrijving"] = merged_products_df["Beschrijving"].str.strip()
    merged_products_df = merged_products_df[merged_products_df["Naam"] != ""]

    # ✅ Keep this line! You still need to save the merged results for use later
    st.session_state["product_push_table"] = merged_products_df

    st.markdown("---")


    # --- Main button ---
    if st.button(t("generate_button")):
    # --- Handle analytics uploads (optional) ---
        if analytics_files:
            video_df, overview_df, audience_df = load_analytics(analytics_files)
            if video_df.empty:
                st.warning(t("err_no_videos"))
                video_df = pd.DataFrame()
                overview_df = pd.DataFrame()
                audience_df = pd.DataFrame()
        else:
            # No uploads — use empty placeholders
            video_df = pd.DataFrame()
            overview_df = pd.DataFrame()
            audience_df = pd.DataFrame()

        norm_video_df = normalise_video_df(video_df)
        brand_cfg = parse_brandbook(brandbook_file, brand)

        # --- Extract performing tags from own data ---
        performing_tags = []
        performing_tag_names = []

        if not video_df.empty:
            performing_tags = extract_performing_hashtags(norm_video_df, top_n=10)
            performing_tag_names = [t[0] for t in performing_tags]

        live_hashtags: List[str] = []
        live_sounds: List[Dict[str, str]] = []


        # The performing hashtags have already been displayed above; avoid duplicate display.

        # --- Fetch Apify trends automatically (unless a manual trend is given) ---
        from typing import Any, Dict, List

        hashtags_data: List[Any] = []
        sounds_data: List[Dict[str, Any]] = []
        live_hashtags: List[Any] = []
        live_sounds: List[Dict[str, Any]] = []

        if main_trend_focus and main_trend_focus.strip():
            # User gave a manual main trend → skip Apify fetch, keep datasets empty
            st.info(
                f"Handmatige hoofdtrend ingevuld: **{main_trend_focus.strip()}**. "
                "Apify-trends worden genegeerd."
            )
        else:
            # Localized spinner message
            with st.spinner(f"{t('fetching_trends_message')} {region}..."):
                hashtags_data, sounds_data = fetch_apify_trends(region, time_range=selected_time_range)
            live_hashtags = hashtags_data or []
            live_sounds = sounds_data or []



        # --- Filter out non-Latin hashtags (Cyrillic, Arabic, etc.) ---

        cleaned_hashtags = []

        for tag in hashtags_data:
           # Extract the text if Apify returned a dict
            if isinstance(tag, dict):
                name = (
                    tag.get("name")
                    or tag.get("hashtagName")
                    or tag.get("title")
                    or tag.get("text")
                    or ""
                )
            else:
                name = str(tag)

            name = name.strip().lstrip("#")

            if is_latin_hashtag(name):
                cleaned_hashtags.append(name)

        live_hashtags = cleaned_hashtags


        print(f"✅ Filtered to {len(hashtags_data)} Latin hashtags (from {len(cleaned_hashtags)}).")



        if len(hashtags_data) < 5:
            st.warning("⚠️ Let op: weinig NL/EN hashtags gevonden na filteren van niet-Latijnse woorden.")


        # --- Show performing hashtags ---
        st.subheader(t("performing_tags"))
        if performing_tags:
            df_tags = pd.DataFrame(performing_tags, columns=["Hashtag", "Gemiddelde successcore"])
            # Localize the column name for average success score
            if get_lang() == "en":
                df_tags = df_tags.rename(columns={"Gemiddelde successcore": t("average_success_score")})
            st.dataframe(df_tags)
        else:
            st.info("Geen presterende hashtags geanalyseerd uit eigen kanaal." if get_lang() == "nl" else "No top-performing hashtags analysed for brand channel.")

        # --- Show live trends (only when no manual main trend is given) ---
        if not main_trend_focus or not main_trend_focus.strip():
            st.subheader(t("live_hashtags"))
            if live_hashtags:
                st.dataframe(pd.DataFrame({"Hashtag": live_hashtags}))
            else:
                st.info("Geen live hashtags gevonden." if get_lang() == "nl" else "No live hashtags found.")

            st.subheader(t("live_sounds"))
            if live_sounds:
                # Build a localized DataFrame for sounds
                df_sounds = pd.DataFrame([
                    {
                        t("sound"): s.get("soundName", ""),
                        t("artist"): s.get("artist", t("unknown_artist")),
                        t("commercial_allowed"): t("yes") if s.get("isCommercial") else t("no")
                    }
                    for s in live_sounds
                ])
                st.dataframe(df_sounds)
            else:
                st.info("Geen live sounds gevonden." if get_lang() == "nl" else "No live sounds found.")
        else:
            st.subheader(t("main_trend_focus"))
            st.info(f"De hoofdtrend **{main_trend_focus.strip()}** is handmatig opgegeven; "
                    "Apify-trends worden niet getoond.")



        # --- Use stub sounds if no live sounds available ---
        if not live_sounds:
            live_sounds = [
                {"soundName": "Upbeat Pop Beat", "isCommercial": True},
                {"soundName": "Calm Guitar Loop", "isCommercial": True},
                {"soundName": "Viral TikTok Sound", "isCommercial": False},
            ]

        # --- Determine trending hashtags to include in the prompt ---
        # Filter live hashtags for relevancy to the niche and remove banned topics
        filtered_trending_tags = filter_trending_hashtags(live_hashtags, niche)
        if filtered_trending_tags:
            trending_for_prompt = filtered_trending_tags[:10]
        else:
            trending_for_prompt = performing_tag_names[:10]

        # --- Generate new ideas using OpenAI ---
        # Ensure the brand name is included in the brand configuration for the prompt
        brand_cfg["brand"] = brand

        # Retrieve the merged products and platform choice from session_state
        product_push_table = st.session_state.get(
            "product_push_table", pd.DataFrame(columns=["Naam", "Beschrijving"])
        )
        # Retrieve platform choice from session state, defaulting to localized "no preference"
        platform_choice = st.session_state.get("platform_choice", t("no_preference"))

        ideas_df = generate_ideas_openai(
            norm_video_df,
            trending_for_prompt,
            live_sounds,
            brand_cfg,
            n_ideas=5,
            niche=niche,
            product_push_table=product_push_table,  # <--- updated variable
            cohort_focus=cohort_focus,
            main_trend_focus=main_trend_focus,
            low_budget=low_budget,
            lang=get_lang(),
        )


        # --- Clean GPT outputs for formatting (remove literal \n, \t) ---
        if ideas_df is not None and not ideas_df.empty:
            for col in ideas_df.columns:
                if ideas_df[col].dtype == "object":
                    ideas_df[col] = ideas_df[col].apply(clean_output_text)


        # --- Display the generated ideas ---
        if ideas_df is not None and not ideas_df.empty:
            st.subheader(t("ideas_header"))
            st.dataframe(ideas_df)


        # --- Export options for the generated ideas ---
        col1, col2 = st.columns(2)

        # ---- CSV Download ----
        with col1:
            csv_data = ideas_df.to_csv(index=False).encode("utf-8")
            st.download_button(
                label=t("download_csv"),
                data=csv_data,
                file_name="tiktok_ideas.csv",
                mime="text/csv",
            )

        # ---- DOCX Download ----
        with col2:
            brand_name = st.session_state.get("brand_name", "Merknaam")
            docx_filename = "tiktok_ideas.docx"
            path = export_ideas_to_docx(ideas_df, brand_name, output_path=docx_filename, lang=get_lang())

            with open(path, "rb") as f:
                st.download_button(
                    label=t("download_docx"),
                    data=f,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_ideas_to_docx(df, brand_name, output_path="tiktok_ideas.docx", lang: str = "nl"):
    """
    Export generated TikTok ideas to a formatted Word document,
    including trend name and explanation ("Waarom actueel").
    """
    doc = Document()

    # --- Loop through ideas ---
    for i, row in df.iterrows():
        # Choose title in either column ('Idee'/'Titel' are your existing field names)
        _title_fallback = row.get('Idee', row.get('Titel', 'Untitled' if lang == "en" else 'Ongetiteld'))
        doc.add_heading(t_format("docx_title", i=i+1, title=_title_fallback), level=1)


        if row.get("Trend"):
            p = doc.add_paragraph()
            p.add_run(t("docx_trend")).bold = True
            p.add_run(str(row["Trend"]))

        if row.get("Waarom actueel"):
            p = doc.add_paragraph()
            p.add_run(t("docx_why_now")).bold = True
            p.add_run(str(row["Waarom actueel"]))

        if row.get("Hook"):
            p = doc.add_paragraph()
            p.add_run(t("docx_hook")).bold = True
            p.add_run(str(row["Hook"]))

        # Caption field (optional)
        if row.get("Caption"):
            p = doc.add_paragraph()
            p.add_run(t("docx_caption")).bold = True
            p.add_run(str(row["Caption"]))

        # --- robust script handling ---
        script = row.get("Script", "")
        if isinstance(script, list):
            script = "\n".join(script)
        elif not isinstance(script, str):
            script = str(script)
        if script:
            p = doc.add_paragraph()
            p.add_run(t("docx_script")).bold = True
            for line in script.split("\n"):
                s = doc.add_paragraph(line)
                s.paragraph_format.left_indent = Pt(12)

        if row.get("Filmtips"):
            p = doc.add_paragraph()
            p.add_run(t("docx_tips")).bold = True
            p.add_run(str(row["Filmtips"]))

        if row.get("Hashtags"):
            p = doc.add_paragraph()
            p.add_run(t("docx_hashtags")).bold = True
            hashtags = (
                ", ".join(row["Hashtags"])
                if isinstance(row["Hashtags"], (list, tuple))
                else str(row["Hashtags"])
            )
            p.add_run(hashtags)

        if row.get("Sound"):
            p = doc.add_paragraph()
            p.add_run(t("docx_sound")).bold = True
            sound_text = str(row["Sound"])
            if "Commercial OK" in row:
                sound_text += t_format("docx_sound_ok", ok=row["Commercial OK"])
            p.add_run(sound_text)

        doc.add_page_break()

    # --- Save file ---
    doc.save(output_path)
    return output_path



# --- Run the app ---
if __name__ == "__main__":
    main()
