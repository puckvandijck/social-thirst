"""
Streamlit app for TikTok content idea generation (V3).

V3 introduces live trend scraping via Apify and content generation via OpenAI.

Features:
 - Uses Apify's TikTok Trends Scraper (clockworks/tiktok-trends-scraper) to
   fetch trending hashtags and songs for the selected region and industry. This
   requires the user to set the environment variable `APIFY_API_TOKEN` with
   their Apify API token. If no token is provided, the app falls back to the
   stub trends defined in V2.
 - Generates content ideas using OpenAI's API (GPT-5 or later). The user must
   set `OPENAI_API_KEY` to their OpenAI API key. A fallback heuristic generator
   is provided if no key is found.
 - Integrates vectorization of captions using TF-IDF to provide context to the
   language model and to compute similarity between trends and existing posts.
 - Respects the brand book guidelines (placeholder parser) and includes them in
   the prompt sent to the model.

Before running this app, install the necessary dependencies:

```
pip install streamlit pandas numpy python-docx scikit-learn openai apify-client
```

Set up environment variables in a `.env` file or your shell:

```
export OPENAI_API_KEY="sk-..."
export APIFY_API_TOKEN="apify-api-token"
```

Then run the app with:

```
streamlit run streamlit_app_v3.py
```

Replace the placeholders with your actual API tokens. The app will attempt to
fetch live trends from Apify; if the token is missing, stub data will be used.
"""

import os
import regex as re
from collections import Counter
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt

# Optional external libraries
try:
    from sklearn.feature_extraction.text import TfidfVectorizer  # type: ignore
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

try:
    import openai  # type: ignore
    OPENAI_AVAILABLE = True
except Exception:
    OPENAI_AVAILABLE = False

try:
    from apify_client import ApifyClient  # type: ignore
    APIFY_AVAILABLE = True
except Exception:
    APIFY_AVAILABLE = False


# -------------------------------
# Helper functions (reused from V2)
# -------------------------------

import regex as re
import unicodedata

def is_latin_hashtag(tag: str) -> bool:
    """
    Return True if the hashtag text is primarily Latin-based
    (Dutch, English, German, etc.) and not Cyrillic, Arabic, or Asian.
    """

    if not tag:
        return False

    # Remove hashtag symbol and trim
    tag = tag.lstrip("#").strip()

    # Normalize and remove invisible Unicode marks
    tag = unicodedata.normalize("NFKC", tag)
    tag = re.sub(r"[\u200b-\u200f\u202a-\u202e\u2060-\u206f]", "", tag)

    # Reject if non-Latin scripts are present
    if re.search(r'[\p{IsCyrillic}\p{IsArabic}\p{IsHan}\p{IsHiragana}\p{IsKatakana}]', tag):
        return False

    # Keep if contains at least one Latin letter
    return bool(re.search(r'[A-Za-zÀ-ÿ]', tag))

def filter_trending_hashtags(tags: List[str], niche: str) -> List[str]:
    """
    Filter trending hashtags for relevancy to the given niche and remove banned or controversial topics.
    Only tags that relate to the niche (e.g. hair care, body care, beauty) or generic event-related words
    are kept. Tags containing political, election or sports-related keywords are excluded.
    """
    banned_keywords = [
        "politic", "politiek", "election", "vote", "voetbal", "soccer", "football",
        "war", "conflict", "controversy", "verkiezing", "verkiezingen"
    ]
    niche_keywords_map = {
        "Hair Care": [
            "hair", "haar", "hairstyle", "haarstijl", "haircare", "haarverzorging",
            "beauty", "glow", "party", "feest", "festival", "dance", "event"
        ],
        "Body Care": [
            "body", "skin", "huid", "bodycare", "body care", "lotion", "moisturizer",
            "wellness", "beauty"
        ],
        "Beauty": [
            "beauty", "glow", "makeup", "style", "fashion", "party", "fest", "festival", "event"
        ],
    }
    filtered: List[str] = []
    for tag in tags:
        tag_name = str(tag).lower()
        # Skip banned or controversial topics
        if any(bad in tag_name for bad in banned_keywords):
            continue
        allowed_keywords = niche_keywords_map.get(niche, [])
        # Keep tag if it contains any niche keyword
        if any(keyword in tag_name for keyword in allowed_keywords):
            filtered.append(tag)
        else:
            # Keep general event/festival/party tags if not explicitly banned
            if any(word in tag_name for word in ["event", "festival", "party", "feest", "fest"]):
                filtered.append(tag)
    return filtered


def detect_file_type(df: pd.DataFrame) -> str:
    cols = set(c.lower().strip() for c in df.columns)
    if {"video title", "video views", "video link"}.issubset(cols):
        return "video"
    if {"profile views", "video views", "reached audience"}.issubset(cols):
        return "overview"
    if {"new followers", "total followers", "engaged audience"}.issubset(cols):
        return "audience"
    return "unknown"


def load_analytics(files: List) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    video_df = pd.DataFrame()
    overview_df = pd.DataFrame()
    audience_df = pd.DataFrame()
    for uploaded_file in files:
        if uploaded_file is None:
            continue
        try:
            df = pd.read_csv(uploaded_file)
        except Exception:
            try:
                df = pd.read_excel(uploaded_file)
            except Exception:
                st.warning(f"Unable to read {uploaded_file.name} as CSV or Excel.")
                continue
        ftype = detect_file_type(df)
        if ftype == "video":
            video_df = df
        elif ftype == "overview":
            overview_df = df
        elif ftype == "audience":
            audience_df = df
    return video_df, overview_df, audience_df


def normalise_video_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [c.strip() for c in df.columns]
    if "Post time" in df.columns:
        df["Post time"] = pd.to_datetime(df["Post time"], errors="coerce")
        df["Date"] = df["Post time"].dt.date
    for col in ["Video views", "Likes", "Comments", "Shares", "Add to Favorites"]:
        if col not in df.columns:
            df[col] = 0
        df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)
    df["Engagement"] = df[["Likes", "Comments", "Shares", "Add to Favorites"]].sum(axis=1)
    df["Engagement rate"] = df["Engagement"] / df["Video views"].replace(0, np.nan)
    df["Engagement rate"] = df["Engagement rate"].fillna(0)
    if not df["Engagement rate"].empty:
        max_rate = df["Engagement rate"].max()
        if max_rate > 0:
            df["Success score"] = df["Engagement rate"] / max_rate
        else:
            df["Success score"] = 0
    else:
        df["Success score"] = 0
    return df


def extract_performing_hashtags(video_df: pd.DataFrame, top_n: int = 10) -> List[Tuple[str, float]]:
    if video_df.empty or "Success score" not in video_df.columns:
        return []
    hashtag_pattern = re.compile(r"#(\w+)")
    hashtag_scores: Dict[str, List[float]] = {}
    for _, row in video_df.iterrows():
        title = row.get("Video title", "")
        success = float(row.get("Success score", 0))
        if isinstance(title, str):
            tags = hashtag_pattern.findall(title)
            for tag in tags:
                tag_l = tag.lower()
                hashtag_scores.setdefault(tag_l, []).append(success)
    averages: Dict[str, float] = {
        tag: (sum(scores) / len(scores) if scores else 0)
        for tag, scores in hashtag_scores.items()
    }
    sorted_tags = sorted(averages.items(), key=lambda x: x[1], reverse=True)
    return sorted_tags[:top_n]


def parse_brandbook_placeholder() -> Dict[str, List[str]]:
    return {
        "tone": "neutraal, vriendelijk, deskundig",
        "avoid": ["garantie", "geneest", "wondermiddel"],
        "must": ["Andrélon", "natuurlijk", "haarverzorging"],
    }


def vectorize_captions(video_df: pd.DataFrame) -> Tuple[Optional[TfidfVectorizer], Optional[np.ndarray]]:
    if not SKLEARN_AVAILABLE:
        return None, None
    titles = [str(t) if not pd.isna(t) else "" for t in video_df.get("Video title", [])]
    if not titles:
        return None, None
    vectorizer = TfidfVectorizer(max_features=200)
    matrix = vectorizer.fit_transform(titles)
    return vectorizer, matrix

from typing import List, Dict, Any
import json

def analyse_trends_meanings(
    client,
    trends: List[str],
    brand_name: str,
    niche: str,
    max_trends: int = 12,
    model: str = "gpt-5.1-mini",
) -> Dict[str, str]:
    """
    Calls the model once to produce short Dutch explanations for each trend.
    Returns a dict {trend: explanation}. Safe for empty/noisy inputs.
    """
    trends = [t for t in (trends or []) if isinstance(t, str) and t.strip()]
    if not trends:
        return {}

    # Deduplicate while preserving order; limit to avoid token bloat
    seen = set()
    clean = []
    for t in trends:
        key = t.strip()
        if key not in seen:
            seen.add(key)
            clean.append(key)
        if len(clean) >= max_trends:
            break

    # Ask for strict JSON to keep parsing robust
    user_prompt = f"""
Je bent een Nederlandse TikTok-trendanalist voor de {niche}-niche (merk: {brand_name}).
Geef per trend een KORTE, praktische uitleg in maximaal 2 zinnen:
- Wat houdt de trend in?
- Hoe manifesteert deze zich in TikTok-content (stijl/format/doelgroep)?

Antwoord ALLEEN als strikt JSON-object waarbij de keys exact de aangeleverde trendnamen zijn, bijv.:
{{
  "clean girl aesthetic": "Korte uitleg...",
  "get ready with me": "Korte uitleg..."
}}

Trends:
{json.dumps(clean, ensure_ascii=False)}
"""

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "Je bent een beknopte, nauwkeurige assisent die uitsluitend valide JSON terugstuurt."},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
        )
        raw = resp.choices[0].message.content.strip()
        # Try to extract JSON even if the model wrapped it in code fences
        if raw.startswith("```"):
            raw = raw.strip("`")
            # remove possible language hints like ```json
            first_newline = raw.find("\n")
            if first_newline != -1:
                raw = raw[first_newline+1:].strip()
        parsed = json.loads(raw)
        # Keep only requested trends; coerce to str
        out: Dict[str, str] = {}
        for t in clean:
            val = parsed.get(t, "")
            if isinstance(val, str):
                out[t] = val.strip()
            else:
                out[t] = str(val)
        return out
    except Exception:
        # Fail-soft: return empty dict on any parsing/network issues
        return {}

# --- Helper: Clean escaped newlines and tabs from GPT output ---
def clean_output_text(text: str) -> str:
    """Replace escaped newlines and stray backslashes with real newlines/tabs."""
    if not isinstance(text, str):
        return text
    return text.replace("\\n", "\n").replace("\\t", "\t").strip()


# -------------------------------
# Live trend scraping via Apify (robust parser)
# -------------------------------
def fetch_apify_trends(
    region: str,
    industry: str = "Beauty & Personal Care",
    results_per_page: int = 50,
    time_range: str = "7"   # ✅ new parameter
) -> Tuple[List[str], List[Dict[str, str]]]:

    """
    Returns:
      - hashtags: List[str]
      - sounds: List[{"soundName": str, "artist": str, "isCommercial": bool}]
    """
    if not APIFY_AVAILABLE:
        st.warning("Apify not available in this environment.")
        return [], []

    token = os.environ.get("APIFY_API_TOKEN")
    if not token:
        st.warning("No Apify token found. Please export APIFY_API_TOKEN.")
        return [], []

    region_to_code = {
        "Nederland": "NL",
        "België": "BE",
        "Duitsland": "DE",
        "Verenigde Staten": "US",
    }
    country_code = region_to_code.get(region, "NL")

    client = ApifyClient(token)
    run_input = {
        "adsCountryCode": country_code,
        "adsApprovedForBusinessUse": True,   # keep True to align with brand-safe music
        "adsScrapeHashtags": True,
        "adsScrapeSounds": True,
        "adsScrapeCreators": False,
        "adsScrapeVideos": False,
        "adsRankType": "popular",
        "adsTimeRange": time_range,   # ✅ dynamic timeframe from user selection,
        "resultsPerPage": results_per_page,
    }

    try:
        run = client.actor("clockworks/tiktok-trends-scraper").call(run_input=run_input, timeout_secs=1200)
    except Exception as e:
        st.error(f"Error running Apify actor: {e}")
        return [], []

    dataset_id = run.get("defaultDatasetId")
    if not dataset_id:
        st.error("No dataset returned from Apify run.")
        return [], []

    hashtags: List[str] = []
    sounds: List[Dict[str, str]] = []

    try:
        for item in client.dataset(dataset_id).iterate_items():
            # Some actors return dicts, never strings — but be defensive:
            if not isinstance(item, dict):
                continue

            itype = str(item.get("type", "")).lower()

            # --- Hashtags ---
            if "hashtag" in itype:
                name = item.get("hashtagName") or item.get("title") or item.get("name")
                if name and name not in hashtags:
                    hashtags.append(name)
                continue

            # --- Sounds (sometimes 'sound', sometimes 'music') ---
            if ("sound" in itype) or ("music" in itype):
                sound_name = item.get("soundName") or item.get("musicName") or item.get("title") or item.get("name")
                artist = item.get("authorName") or item.get("artistName") or item.get("creatorName") or "Onbekend"
                commercial = item.get("isCommercial")
                if commercial is None:
                    commercial = item.get("commercial_ok", item.get("commercial", False))
                if sound_name:
                    sounds.append({
                        "soundName": sound_name,
                        "artist": artist,
                        "isCommercial": bool(commercial),
                    })
                continue

            # Optional: heuristic fallback (rare)
            url = item.get("url", "")
            if url and "/tag/" in url:
                name = item.get("title") or item.get("name")
                if name and name not in hashtags:
                    hashtags.append(name)

    except Exception as e:
        st.error(f"Error parsing Apify dataset: {e}")
        return [], []

    if not hashtags:
        st.warning("Geen live hashtags gevonden (controleer filters of land).")
    if not sounds:
        st.warning("Geen live sounds gevonden (controleer filters of land).")

    return hashtags, sounds



# -------------------------------
# Idea generation with OpenAI (improved V2 — trend-weighted & diverse)
# -------------------------------

from openai import OpenAI
import os, json
from datetime import datetime
import pandas as pd
import streamlit as st

def generate_ideas_openai(
    video_df: pd.DataFrame,
    trending_tags: List[str],
    sounds: List[Dict[str, str]],
    brand_cfg: Dict[str, List[str]],
    n_ideas: int = 5,
    niche: str = "Beauty",
    product_push_table: pd.DataFrame | None = None,
    cohort_focus: str = "",
    main_trend_focus: str = "",
) -> pd.DataFrame:
    """Generate diverse, trend-weighted TikTok ideas using OpenAI GPT-5."""
    trend_analyses = {}  # ✅ ensures variable exists in all cases
    
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.warning("OPENAI_API_KEY is niet ingesteld.")

    client = OpenAI(api_key=api_key)

    # --- Context preparation ---
    # Use only the top 3 videos to prevent over-indexing on historical performance
    top_videos = video_df.sort_values(by="Success score", ascending=False).head(3)
    video_summaries = [
        {
            "title": str(row.get("Video title", "")),
            "success_score": float(row.get("Success score", 0)),
            "engagement_rate": float(row.get("Engagement rate", 0)),
        }
        for _, row in top_videos.iterrows()
    ]

    tone = brand_cfg.get("tone", "")
    avoid_words = ", ".join(brand_cfg.get("avoid", []))
    must_words = ", ".join(brand_cfg.get("must", []))
    today = datetime.now().strftime("%d %B %Y")

    # --- Construct creative, trend-prioritizing prompt ---
    product_block = ""
    try:
        if product_push_table is not None and not product_push_table.empty:
            tbl = product_push_table.copy()
            # Normalize column names
            tbl.columns = [str(c).strip() for c in tbl.columns]

            # Ensure required columns exist
            for col in ["Productnaam", "Beschrijving", "USP's"]:
                if col not in tbl.columns:
                    tbl[col] = ""

            # Replace NaN with empty strings and cast to str for safety
            tbl = tbl.fillna("").astype(str)

            product_rows = []
            for _, row in tbl.iterrows():
                name = row["Productnaam"].strip()
                desc = row["Beschrijving"].strip()
                usps = row["USP's"].strip()

                if name or desc or usps:
                    line = f"- **{name or 'Onbenoemd product'}**"
                    if desc:
                        line += f": {desc}"
                    if usps:
                        line += f" | USP’s: {usps}"
                    product_rows.append(line)

            if product_rows:
                product_block = "\n### Product of range om te pushen\n" + "\n".join(product_rows) + "\n"
    except Exception:
        # Fail-safe: if anything goes wrong, just omit the product block
        product_block = ""

    # --- Determine which trends to include in the prompt + analyse meanings ---
    trend_section_text = "### Trends\nGeen trendinformatie beschikbaar."

    if main_trend_focus and main_trend_focus.strip():
        # Manual main trend → ignore Apify & analytics trends
        focus_raw = main_trend_focus.strip()
        focus_nohash = focus_raw.lstrip("#")
        # Show BOTH the hashtag and raw term to anchor the exact token in the model
        top_hashtags = [f"#{focus_nohash}", focus_raw]
        sounds_for_prompt = []
        trends_to_analyse = [focus_raw]

        brand_name_for_ai = brand_cfg.get("brand", "onbekend")
        trend_analyses = analyse_trends_meanings(
            client=client,
            trends=trends_to_analyse,
            brand_name=brand_name_for_ai,
            niche=niche,
            max_trends=12,
            model="gpt-5.1-mini",
        )
    
    analyses_markdown = ""
    if trend_analyses:
        bullets = []
        for t in trends_to_analyse:
            expl = trend_analyses.get(t, "").strip()
            bullets.append(f"- **{t}** — {expl}" if expl else f"- **{t}**")
        analyses_markdown = "\n\n### Korte trendanalyses\n" + "\n".join(bullets)

    # --- FIXED STRUCTURE ---
    if main_trend_focus and main_trend_focus.strip():
        trend_section_text = (
            "### Hoofdtrend om op te focussen\n"
            f"{focus_raw}\n"
            "### Verplichte hashtag/term\n"
            f"#{focus_nohash} (en/of '{focus_raw}')\n"
            "### Instructie\n"
            f"Behandel deze hoofdtrend als bewezen relevant voor {brand_cfg.get('brand','het merk')}. "
            "Integreer deze hoofdtrend in elk idee en gebruik de exacte tag/term in de hashtag-lijst."
            f"{analyses_markdown}"
        )

    else:
        # Build list of allowed trending hashtags from 'trending_tags' parameter
        allowed_tags = []
        for t in (trending_tags or []):
            if not t:
                continue
            tag = t if t.startswith("#") else f"#{t}"
            if tag not in allowed_tags:
                allowed_tags.append(tag)

        if allowed_tags:
            bullet_list = "\n".join(f"- {t}" for t in allowed_tags[:10])
            trend_section_text = (
                "### Toegestane trending hashtags (EXACTE strings)\n"
                f"{bullet_list}\n"
                "### Instructie\n"
                "Kies per idee precies één van de bovenstaande hashtags als het veld 'trend_used'. "
                "Neem die tag ongewijzigd (exacte string) op in 'hashtags'. "
                "Pas de tag NIET aan: geen extra woorden, geen meervoud, geen prefix/suffix, geen varianten."
            )
        else:
            trend_section_text = "### Trends\nGeen trendinformatie beschikbaar."



    #### prompt ####

    prompt = f"""
    Je bent een creatieve TikTok-strateeg gespecialiseerd in Gen Z-content en merkactivaties voor het merk {brand_cfg.get('brand', 'onbekend')} in de {niche}-niche. 
    Gebruik de volgende informatie om {n_ideas} nieuwe, actuele en creatieve TikTok-video-ideeën te bedenken.
    

    ### Informatie

    # Merktoon
    {tone or "Professioneel maar speels"}
    Moet bevatten: {must_words or "geen specifieke richtlijnen"}
    Vermijd: {avoid_words or "geen specifieke verboden woorden"}

    # Product/range om te verwerken in content idee
    {product_block if product_block else ""}

    # Toppresterende video's (context, niet dominant)
    {video_summaries}

    # Trending hashtags (dominant)
    {trend_section_text}


    ### Opdracht
    1. Analyseer welke trends het meest relevant zijn voor {brand_cfg.get('brand', 'het merk')}.
    2. Bedenk vervolgens {n_ideas} creatieve TikTok-ideeën die inspelen op deze trends. De trends zijn tijdelijk en belangrijker dan historische data — gebruik ze actief in je ideeën.
    3. Zorg dat de ideeën variëren in type (educatief, humor, transformatie, lifestyle, storytelling, enz.), en vermijdt reclamejargon.
    4. {(
      "Gebruik uitsluitend de opgegeven hoofdtrend. "
      "Voeg de exacte term toe aan de hashtaglijst van ieder idee: "
      f"#{focus_nohash} (en/of '{focus_raw}'). Negeer overige trends."
      if main_trend_focus and main_trend_focus.strip()
      else "Integreer minimaal één actuele trend (hashtag of sound) per idee en leg de koppeling kort uit."
    )}
    5. Gebruik de merktoon, maar geef ruimte aan creatieve vrijheid.
    6. {(
          f"Focus op het cohort **{cohort_focus.strip()}** en zorg dat er minstens één product uit de range **{', '.join(product_push_table['Productnaam'].dropna().unique())}** wordt meegenomen."
          if cohort_focus and product_push_table is not None and not product_push_table.empty
          else f"Focus op het cohort **{cohort_focus.strip()}**."
          if cohort_focus
          else f"Zorg dat er minstens één product uit de range **{', '.join(product_push_table['Productnaam'].dropna().unique())}** wordt meegenomen."
          if product_push_table is not None and not product_push_table.empty
          else ""
        )}
    7. Beschrijf per idee een creatief “Script” in 4–6 korte shots (SHOT 1, SHOT 2, …) met voice-over, on-screen text of handelingen. 
       Zorg dat de shots logisch op elkaar volgen binnen één herkenbare setting of een natuurlijk verloop. Geen abrupte wissels zonder context, zorg dat de flow duidelijk is. 
       Zorg dat het merk en product natuurlijk in het verhaal passen.
    8. Vermijd politieke of controversiële onderwerpen en zorg dat de trends logisch passen bij de niche. 
       Bijvoorbeeld: voor haarverzorging zijn feesten en evenementen relevant (zoals festivals), maar sportwedstrijden 
       of politieke kwesties niet.

    BELANGRIJKE REGELS VOOR TRENDS EN HASHTAGS:
    - Gebruik een trending hashtag alleen als deze duidelijk aansluit bij het merk of de niche.
    - Kies per idee één trending hashtag en zet die exact (ongewijzigd, inclusief #) in 'trend_used'.
    - Neem diezelfde hashtag ook letterlijk op in 'hashtags'.
    - 'hashtags' bevat 3–5 items die allemaal met # beginnen.
    - Geef uitsluitend valide JSON terug met de velden: title, hook, script, filming_tips, hashtags, sound, commercial_ok, short_explanation, trend_used.


    """

    # --- GPT call ---
    try:
        response = client.chat.completions.create(
            model="gpt-5",  # or "gpt-4o" if preferred
            messages=[
                {"role": "system", "content": "Je bent een ervaren TikTok-marketeer en trendanalist."},
                {"role": "user", "content": prompt},
            ],
            temperature=1.0,  # increase creativity and variation
        )

    except Exception as e:
        st.error(f"❌ Fout bij het aanroepen van OpenAI: {e}")
        return pd.DataFrame()  # stop here if the API call failed

    # continue normally if no error
    records = []
    allowed_tags_set = set()
    # Reconstruct the allowed set so we can validate, based on what the prompt showed:
    if main_trend_focus and main_trend_focus.strip():
        focus_raw = main_trend_focus.strip()
        focus_tag = focus_raw if focus_raw.startswith("#") else f"#{focus_raw.lstrip('#')}"
        allowed_tags_set.add(focus_tag.lower())
    else:
        for t in (trending_tags or []):
            if not t:
                continue
            tag = t if t.startswith("#") else f"#{t}"
            allowed_tags_set.add(tag.lower())

    # --- Parse model reply safely ---
    try:
        reply = response.choices[0].message.content.strip()
        if reply.startswith("```"):
            reply = reply.strip("`").split("json", 1)[-1].strip()
        ideas_list = json.loads(reply)
    except Exception as e:
        st.warning(f"⚠️ Fout bij het lezen van modeloutput: {e}")
        ideas_list = []

    # continue
    for idea in ideas_list[:n_ideas]:
        title = idea.get("title", "")
        hook = idea.get("hook", "")
        script = idea.get("script", "")

        # filming tips → string
        filming_tips_raw = idea.get("filming_tips", "")
        if isinstance(filming_tips_raw, list):
            filming_tips = "; ".join(map(str, filming_tips_raw))
        else:
            filming_tips = str(filming_tips_raw or "")

        # hashtags → list normalization
        hashtags_raw = idea.get("hashtags", [])
        if isinstance(hashtags_raw, str):
            # split on commas/spaces, keep entries with a leading #
            parts = [h.strip() for h in re.split(r"[,\s]+", hashtags_raw) if h.strip()]
            hashtags = [h if h.startswith("#") else f"#{h.lstrip('#')}" for h in parts]
        elif isinstance(hashtags_raw, list):
            hashtags = [h if str(h).startswith("#") else f"#{str(h).lstrip('#')}" for h in hashtags_raw]
        else:
            hashtags = []

        # trend_used → force exact presence in hashtags if it's valid/allowed
        trend_used = idea.get("trend_used", "")
        if trend_used:
            trend_used = trend_used if trend_used.startswith("#") else f"#{trend_used.lstrip('#')}"
            # If we have an allowed set, prefer only exact allowed tags
            if allowed_tags_set and trend_used.lower() not in allowed_tags_set:
                # if model invented/modified a tag, try to fallback to first allowed tag
                trend_used = next(iter(allowed_tags_set)).lower()
                # ensure it has '#'
                if not str(trend_used).startswith("#"):
                    trend_used = f"#{str(trend_used).lstrip('#')}"
            # ensure trend_used is included in hashtags (exact string)
            if trend_used not in hashtags:
                hashtags = [trend_used] + [h for h in hashtags if h.lower() != trend_used.lower()]
        else:
            # No trend_used provided: if allowed set exists, insert the first allowed tag
            if allowed_tags_set:
                trend_used = next(iter(allowed_tags_set))
                if not str(trend_used).startswith("#"):
                    trend_used = f"#{str(trend_used).lstrip('#')}"
                if trend_used not in hashtags:
                    hashtags = [trend_used] + hashtags

        # sound & flags
        sound = idea.get("sound", "")
        commercial_ok = "Ja" if idea.get("commercial_ok", True) else "Nee"
        short_explanation = idea.get("short_explanation", "")

        records.append({
            "Titel": title,
            "Hook": hook,
            "Script": script,
            "Filming tips": filming_tips,
            "Hashtags": ", ".join(dict.fromkeys(hashtags)),  # keep order, dedupe
            "Sound": sound,
            "Commercial OK": commercial_ok,
            "Waarom actueel": short_explanation,
            "Trend": trend_used or "",  # expose for later export/use
        })

    return pd.DataFrame(records)


# -------------------------------
# Brand book PDF parser (improved robust version)
# -------------------------------
import PyPDF2
import regex as re
import json
import ast
from openai import OpenAI

def clean_pdf_text(text: str) -> str:
    """Normalize extracted PDF text for better GPT comprehension."""
    text = re.sub(r"\s{2,}", " ", text)  # collapse multiple spaces
    text = re.sub(r"\n{2,}", "\n", text)  # reduce multiple newlines
    text = re.sub(r"[^a-zA-Z0-9.,;:!?()&%€$'\-\n ]", "", text)  # strip weird symbols
    text = re.sub(r"(?i)page\s*\d+", "", text)  # remove page markers
    text = text.strip()
    return text

def parse_brandbook(brandbook_file, brand_name: str = "Onbekend") -> Dict[str, List[str]]:
    """Extract tone, must/avoid words, and brand essence from a brand book PDF using GPT."""
    if not brandbook_file:
        st.info("Geen brand book geüpload. Gebruik standaard merktoon.")
        return {"tone": "", "must": [], "avoid": []}

    # --- Step 1: Read and clean PDF text ---
    pdf_reader = PyPDF2.PdfReader(brandbook_file)
    text = ""
    for page in pdf_reader.pages:
        try:
            text += page.extract_text() or ""
        except Exception:
            continue

    cleaned_text = clean_pdf_text(text)

    if len(cleaned_text) < 500:
        st.warning("PDF bevat weinig leesbare tekst — mogelijk is het document grotendeels beeldmateriaal.")
        return {"tone": "", "must": [], "avoid": []}

    # --- Step 2: Analyze using GPT ---
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.warning("OPENAI_API_KEY is niet ingesteld. Kan brand book niet analyseren.")
        return {"tone": "", "must": [], "avoid": []}

    client = OpenAI(api_key=api_key)

    prompt = f"""
    Je bent een Nederlandse merkstrateeg. Analyseer de onderstaande tekst uit een merkboek van {brand_name}.
    Beschrijf:
    1. De toon van stem (max 40 woorden)
    2. 3–5 kernwoorden of thema’s die de merkidentiteit definiëren ("must")
    3. 3–5 woorden of concepten die de merktoon zouden verstoren ("avoid")

    Geef het antwoord als geldige JSON met velden: tone, must, avoid.

    MERKBOEK TEKST:
    {cleaned_text}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-5",  # or "gpt-4o" depending on your setup
            messages=[
                {"role": "system", "content": "Je bent een ervaren Nederlandse merkstrateeg."},
                {"role": "user", "content": prompt},
            ],
        )
        reply = response.choices[0].message.content.strip()

        try:
            cfg = json.loads(reply)
        except Exception:
            cfg = ast.literal_eval(reply)

        # Validate JSON structure
        if not isinstance(cfg, dict):
            raise ValueError("Invalid JSON structure returned by GPT")

        # Ensure required keys exist
        for key in ["tone", "must", "avoid"]:
            if key not in cfg:
                cfg[key] = [] if key != "tone" else ""

        return cfg

    except Exception as e:
        st.error(f"Fout bij het analyseren van brand book: {e}")
        return {"tone": "", "must": [], "avoid": []}

# -------------------------------
# Streamlit UI
# -------------------------------
def main():
    st.set_page_config(
        layout="wide"
    )

    st.title("💜 B&W TikTok Content Agent")

    # --- Custom light-purple theme styling ---
    st.markdown("""
        <style>
            /* Main app background */
            .stApp {
                background-color: #FFFFFF !important;  /* white lavender */
            }
    
            /* Sidebar background and text */
            section[data-testid="stSidebar"] {
                background-color: #EDE7F6;  /* light pastel purple */
                color: #2C1A4D;  /* deep plum text */
            }

            /* Sidebar headings and labels */
            section[data-testid="stSidebar"] h1,
            section[data-testid="stSidebar"] h2,
            section[data-testid="stSidebar"] h3,
            section[data-testid="stSidebar"] p,
            section[data-testid="stSidebar"] label,
            section[data-testid="stSidebar"] span {
                color: #2C1A4D !important;
            }

            /* Sidebar buttons */
            section[data-testid="stSidebar"] button {
                background-color: #6A1B9A !important;  /* strong purple */
                color: #FFFFFF !important;
                border-radius: 8px !important;
            }

            /* Sidebar button hover */
            section[data-testid="stSidebar"] button:hover {
                background-color: #4A0072 !important;
                color: #FFFFFF !important;
            }

            /* Main area buttons */
            button[kind="primary"] {
                background-color: #6A1B9A !important;
                color: #FFFFFF !important;
                border-radius: 8px !important;
            }

            /* Subtle card effect for main content blocks */
            div[data-testid="stVerticalBlock"] > div {
                background-color: #FFFFFF;
                padding: 1.2rem;
                border-radius: 12px;
                box-shadow: 0 2px 8px rgba(106, 27, 154, 0.08); /* soft purple shadow */
                margin-bottom: 1rem;
            }

            /* General text color */
            div[data-testid="stMarkdownContainer"] p {
                color: #2C1A4D;
                font-size: 16px;
            }

            /* Divider line */
            hr {
                border: 1px solid #E0D7F6;
            }

            /* Headers */
            h1, h2, h3 {
                color: #4A0072;
            }
        </style>
    """, unsafe_allow_html=True)


    # --- Sidebar controls ---
       
    with st.sidebar:
        st.header("Instellingen")
        region = st.selectbox(
            "Kies regio",
            ["Nederland", "België", "Duitsland", "Verenigde Staten"],
            index=0,
            help="De regio bepaalt welke live trending data we ophalen."
        )

            # --- Time range selector for trends ---
        time_range = st.selectbox(
            "Selecteer tijdsperiode voor trends",
            ["7 dagen", "30 dagen", "120 dagen"],
            index=0,
            help="Bepaalt over welke periode de trending hashtags en sounds worden opgehaald."
        )

        time_range_map = {
            "7 dagen": "7",
            "30 dagen": "30",
            "120 dagen": "120"
        }
        selected_time_range = time_range_map[time_range]



        # NEW: Brand selection
        brand = st.selectbox(
            "Selecteer merk",
            ["Andrelon", "Dove", "Vaseline"],
            index=0,
            help="Kies het merk waarvoor je contentideeën wilt genereren."
        )

        # Backend mapping of brand → niche
        brand_to_niche = {
            "Andrelon": "Hair Care",
            "Dove": "Body Care",
            "Vaseline": "Body Care",
        }
        niche = brand_to_niche.get(brand, "Beauty")

        ## cohort selection ##
        cohort_focus = st.text_input(
            "Cohort om op te focussen (optioneel)",
            value="",
            help="Bijv. 'Romance Seekers'"
        )


        # --- Optional main trend focus ---
        main_trend_focus = st.text_input(
            "Trend om op te focussen (optioneel)",
            value="",
            help="Bijv. 'group7', 'soft girl aesthetic', 'de-influencing', enz."
        )



    # --- File uploads ---
    st.subheader("Upload Masterbrand Data")
    analytics_files = st.file_uploader(
        "Selecteer één of meerdere CSV/XLSX bestanden", accept_multiple_files=True
    )
    brandbook_file = st.file_uploader(
        "Upload je brand book (PDF, optioneel)", type=["pdf"]
    )

    st.markdown("### 💄 Product of range om te pushen (optioneel)")
    st.write(
        "Voeg hieronder de producten of range toe die je wilt pushen. "
        "Je kunt meerdere rijen toevoegen met naam, beschrijving en USP's. "
        "Laat dit leeg als je geen specifieke producten wilt highlighten."
    )

    # Define an editable table with default empty row
    import pandas as pd
    default_data = pd.DataFrame(
        [{"Productnaam": "", "Beschrijving": "", "USP's": ""}]
    )

    product_push_table = st.data_editor(
        default_data,
        num_rows="dynamic",
        width="stretch",
        key="product_push_table",
    )

    st.markdown("---")

    # --- Main button ---
    if st.button("Genereer ideeën"):
        if not analytics_files:
            st.error("Upload minstens één analytics-bestand.")
            return

        video_df, overview_df, audience_df = load_analytics(analytics_files)
        if video_df.empty:
            st.error("Geen videobestand gevonden in de uploads.")
            return

        norm_video_df = normalise_video_df(video_df)
        brand_cfg = parse_brandbook(brandbook_file, brand)

        # --- Extract performing tags from own data ---
        performing_tags = extract_performing_hashtags(norm_video_df, top_n=10)
        performing_tag_names = [t[0] for t in performing_tags]

        live_hashtags: List[str] = []
        live_sounds: List[Dict[str, str]] = []

        # The performing hashtags have already been displayed above; avoid duplicate display.

        # --- Fetch Apify trends automatically (unless a manual trend is given) ---
        from typing import Any, Dict, List

        hashtags_data: List[Any] = []
        sounds_data: List[Dict[str, Any]] = []
        live_hashtags: List[Any] = []
        live_sounds: List[Dict[str, Any]] = []

        if main_trend_focus and main_trend_focus.strip():
            # User gave a manual main trend → skip Apify fetch, keep datasets empty
            st.info(
                f"Handmatige hoofdtrend ingevuld: **{main_trend_focus.strip()}**. "
                "Apify-trends worden genegeerd."
            )
        else:
            # No manual trend → automatically use live Apify trends
            st.info("Geen hoofdtrend ingevuld — live trends van Apify worden automatisch gebruikt.")
            with st.spinner(f"Bezig met ophalen van TikTok trends voor {region}..."):
                hashtags_data, sounds_data = fetch_apify_trends(region, time_range=selected_time_range)
            live_hashtags = hashtags_data or []
            live_sounds = sounds_data or []



        # --- Filter out non-Latin hashtags (Cyrillic, Arabic, etc.) ---

        cleaned_hashtags = []

        for tag in hashtags_data:
           # Extract the text if Apify returned a dict
            if isinstance(tag, dict):
                name = (
                    tag.get("name")
                    or tag.get("hashtagName")
                    or tag.get("title")
                    or tag.get("text")
                    or ""
                )
            else:
                name = str(tag)

            name = name.strip().lstrip("#")

            if is_latin_hashtag(name):
                cleaned_hashtags.append(name)

        live_hashtags = cleaned_hashtags


        st.info(f"✅ Filtered to {len(hashtags_data)} Latin hashtags (from {len(cleaned_hashtags)}).")



        if len(hashtags_data) < 5:
            st.warning("⚠️ Let op: weinig NL/EN hashtags gevonden na filteren van niet-Latijnse woorden.")


        # --- Show performing hashtags ---
        st.subheader("Presterende hashtags uit je data")
        if performing_tags:
            df_tags = pd.DataFrame(performing_tags, columns=["Hashtag", "Gemiddelde successcore"])
            st.dataframe(df_tags)
        else:
            st.info("Geen presterende hashtags gevonden.")

        # --- Show live trends (only when no manual main trend is given) ---
        if not main_trend_focus or not main_trend_focus.strip():
            st.subheader("Live trending hashtags (Apify)")
            if live_hashtags:
                st.dataframe(pd.DataFrame({"Hashtag": live_hashtags}))
            else:
                st.info("Geen live hashtags gevonden.")

            st.subheader("Live trending sounds (Apify)")
            if live_sounds:
                st.dataframe(pd.DataFrame([
                    {
                        "Sound": s.get("soundName", ""),
                        "Artiest": s.get("artist", "Onbekende artiest"),
                        "Commercieel toegestaan": "Ja" if s.get("isCommercial") else "Nee"
                    }
                    for s in live_sounds
                ]))
            else:
                st.info("Geen live sounds gevonden.")
        else:
            st.subheader("Hoofdtrend om op te focussen")
            st.info(f"De hoofdtrend **{main_trend_focus.strip()}** is handmatig opgegeven; "
                    "Apify-trends worden niet getoond.")



        # --- Use stub sounds if no live sounds available ---
        if not live_sounds:
            live_sounds = [
                {"soundName": "Upbeat Pop Beat", "isCommercial": True},
                {"soundName": "Calm Guitar Loop", "isCommercial": True},
                {"soundName": "Viral TikTok Sound", "isCommercial": False},
            ]

        # --- Determine trending hashtags to include in the prompt ---
        # Filter live hashtags for relevancy to the niche and remove banned topics
        filtered_trending_tags = filter_trending_hashtags(live_hashtags, niche)
        if filtered_trending_tags:
            trending_for_prompt = filtered_trending_tags[:10]
        else:
            trending_for_prompt = performing_tag_names[:10]

        # --- Generate new ideas using OpenAI ---
        # Ensure the brand name is included in the brand configuration for the prompt
        brand_cfg["brand"] = brand
        ideas_df = generate_ideas_openai(
            norm_video_df,
            trending_for_prompt,
            live_sounds,
            brand_cfg,
            n_ideas=5,
            niche=niche,
            product_push_table=product_push_table,
            cohort_focus=cohort_focus,
            main_trend_focus=main_trend_focus,
        )

        # --- Clean GPT outputs for formatting (remove literal \n, \t) ---
        if ideas_df is not None and not ideas_df.empty:
            for col in ideas_df.columns:
                if ideas_df[col].dtype == "object":
                    ideas_df[col] = ideas_df[col].apply(clean_output_text)


        # --- Display the generated ideas ---
        if ideas_df is not None and not ideas_df.empty:
            st.subheader("Gegenereerde contentideeën")
            st.dataframe(ideas_df)


        # --- Export options for the generated ideas ---
        col1, col2 = st.columns(2)

        # ---- CSV Download ----
        with col1:
            csv_data = ideas_df.to_csv(index=False).encode("utf-8")
            st.download_button(
                label="📊 Download als CSV",
                data=csv_data,
                file_name="tiktok_ideas.csv",
                mime="text/csv",
            )

        # ---- DOCX Download ----
        with col2:
            brand_name = st.session_state.get("brand_name", "Merknaam")
            docx_filename = "tiktok_ideas.docx"
            path = export_ideas_to_docx(ideas_df, brand_name, output_path=docx_filename)

            with open(path, "rb") as f:
                st.download_button(
                    label="📄 Download als DOCX",
                    data=f,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_ideas_to_docx(df, brand_name, output_path="tiktok_ideas.docx"):
    """
    Export generated TikTok ideas to a formatted Word document,
    including trend name and explanation ("Waarom actueel").
    """
    doc = Document()

    # --- Loop through ideas ---
    for i, row in df.iterrows():
        doc.add_heading(f"Idee {i+1}: {row.get('Idee', row.get('Titel', 'Ongetiteld'))}", level=1)

        if row.get("Trend"):
            p = doc.add_paragraph()
            p.add_run("Trendfocus: ").bold = True
            p.add_run(str(row["Trend"]))

        if row.get("Waarom actueel"):
            p = doc.add_paragraph()
            p.add_run("Waarom actueel: ").bold = True
            p.add_run(str(row["Waarom actueel"]))

        if row.get("Hook"):
            p = doc.add_paragraph()
            p.add_run("Hook: ").bold = True
            p.add_run(str(row["Hook"]))

        # --- robust script handling ---
        script = row.get("Script", "")
        if isinstance(script, list):
            script = "\n".join(script)
        elif not isinstance(script, str):
            script = str(script)
        if script:
            p = doc.add_paragraph()
            p.add_run("Script:").bold = True
            for line in script.split("\n"):
                s = doc.add_paragraph(line)
                s.paragraph_format.left_indent = Pt(12)

        if row.get("Filmtips"):
            p = doc.add_paragraph()
            p.add_run("Filmtips: ").bold = True
            p.add_run(str(row["Filmtips"]))

        if row.get("Hashtags"):
            p = doc.add_paragraph()
            p.add_run("Hashtags: ").bold = True
            hashtags = (
                ", ".join(row["Hashtags"])
                if isinstance(row["Hashtags"], (list, tuple))
                else str(row["Hashtags"])
            )
            p.add_run(hashtags)

        if row.get("Sound"):
            p = doc.add_paragraph()
            p.add_run("Sound: ").bold = True
            sound_text = str(row["Sound"])
            if "Commercial OK" in row:
                sound_text += f" (Commercieel toegestaan: {row['Commercial OK']})"
            p.add_run(sound_text)

        doc.add_page_break()

    # --- Save file ---
    doc.save(output_path)
    return output_path



# --- Run the app ---
if __name__ == "__main__":
    main()
